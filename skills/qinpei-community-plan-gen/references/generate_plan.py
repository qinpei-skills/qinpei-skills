# -*- coding: utf-8 -*-
# 批量生成多赛道「事业群带教方案」Word（基于 qinpei-community-ops-trust v2.0 方法论）
# 通用版：不绑定任何品牌名（某营养保健品牌等），适用大健康/分享经济团队长
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

GREEN = RGBColor(0x2E, 0x7D, 0x4F)
AMBER = RGBColor(0xB8, 0x6A, 0x1B)
DARK = RGBColor(0x33, 0x33, 0x33)

def set_cn(run, name='微软雅黑'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def build(track, out_dir='/Users/queen1015/Desktop'):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def h1(t):
        p = doc.add_heading(level=1); r = p.add_run(t)
        r.font.color.rgb = GREEN; r.font.size = Pt(15); set_cn(r); return p
    def h2(t):
        p = doc.add_heading(level=2); r = p.add_run(t)
        r.font.color.rgb = AMBER; r.font.size = Pt(12.5); set_cn(r); return p
    def para(t, bold=False, size=10.5, color=DARK, italic=False):
        p = doc.add_paragraph(); r = p.add_run(t)
        r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.italic = italic; set_cn(r); return p
    def bullet(t, pre=None):
        p = doc.add_paragraph(style='List Bullet')
        if pre:
            r = p.add_run(pre); r.bold = True; set_cn(r)
        r = p.add_run(t); set_cn(r); return p
    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ''; rp = hdr[i].paragraphs[0].add_run(h)
            rp.bold = True; rp.font.size = Pt(9.5); set_cn(rp)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ''; rp = cells[i].paragraphs[0].add_run(v)
                rp.font.size = Pt(9.5); set_cn(rp)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows: row.cells[i].width = Inches(w)
        return t

    # ---- 标题（不绑品牌） ----
    ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = ti.add_run('团队长 · {}IP\n事业群带教方案'.format(track['name'])); tr.bold = True
    tr.font.size = Pt(20); tr.font.color.rgb = GREEN; set_cn(tr)
    su = doc.add_paragraph(); su.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = su.add_run('养伙伴 · 带出{}主理人 ｜ 赋能 · 出单 · 发现业务伙伴'.format(track['name']))
    sr.font.size = Pt(11); sr.font.color.rgb = AMBER; set_cn(sr)
    no = doc.add_paragraph(); no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = no.add_run('适用：大健康/分享经济团队长 ｜ 赛道：{} ｜ 群类型：④ 事业群'.format(track['track']))
    nr.font.size = Pt(9); nr.font.color.rgb = DARK; set_cn(nr)
    doc.add_paragraph('─' * 28).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 一、定位与目标 ----
    h1('一、群定位与目标')
    para('这是你的「事业群」——团队长的训练场，不是卖货群，是「养伙伴、养运营伙伴」的地方。', bold=True)
    bullet('周期：3-6 个月，跑得好可延长至一年。', '【生命周期】 ')
    bullet('产品/事业内容占比可高，但也要穿插生活和陪伴，别变成只有喊口号的群。', '【内容占比】 ')
    bullet('目标不是你多能说，是伙伴因为你而成长。', '【价值锚点】 ')
    para('把大目标拆成 4 个可落地的小目标：', bold=True)
    table(['#', '本次核心目标', '怎么衡量做到了'],
          [['1', '赋能伙伴成长', '伙伴能独立建群、排内容、做打卡'],
           ['2', '带出{}主理人'.format(track['name']), '有 3-5 人能独立带{}打卡群/养客群'.format(track['name'])],
           ['3', '卖产品 & 解决方案', '伙伴走通第一单（产品+组合方案）'],
           ['4', '发现潜在业务伙伴', '从活跃伙伴里识别出 2-3 个可深度带教的人']],
          widths=[0.4, 2.2, 3.6])
    para('沁珮关键心法（带教底层逻辑）：', bold=True, color=AMBER)
    bullet('带教 ≠ 培训：教的少而精，动作 > 原理。成年人学习是问题导向、即时有用。')
    bullet('先去做再完美：给伙伴舞台，在实战里长本事，不是先听课。')
    bullet('培养优秀人才是核心：事业群的价值不是你多能说，是伙伴因为你而成长。')

    # ---- 二、带教全链路 ----
    h1('二、带教全链路落地（针对{}主理人）'.format(track['name']))
    para('按 skill 带教全链路 5 步走，每一步都给工具和动作，不空谈。', bold=True)
    h2('第1步：招募 / 邀约 —— 找「事业目标坚定」的人')
    bullet('不是拉人头，是筛人。{}赛道最好的苗子在你的「{}打卡群 / 养客群」里：效果明显 + 爱分享 + 主动问方法的人。'.format(track['name'], track['name']))
    bullet('一对一咨询钩子（核心方法论三 A版）："{}"'.format(track['recruit']))
    bullet('给「内测资格 / 优先名额」特殊感，让 TA 觉得被看见。')
    h2('第2步：培训 / 上手 —— 给工具 + 陪跑第一周')
    bullet('工具包四件套：六大维度 / 黄金四段式 / 排期模板 / 三层信任。', '【给工具】 ')
    bullet('第一周陪跑：一起建群、一起排内容，做给他看，再让他自己做你复盘。', '【陪跑】 ')
    h2('第3步：实战带教 —— 每周固定带教时间 + 话术复盘')
    bullet('每周固定「带教时间」（见第四节模板D 周二）。', '【节奏】 ')
    bullet('把伙伴遇到的卡点变成群内公开课："上周好几个伙伴卡在破冰，今天咱们一起练一句。"', '【复盘】 ')
    h2('第4步：转化 / 出单 —— 帮伙伴走通第一单')
    para('{}版五步（养客→处关系→挖需求→邀客户→给方案）：'.format(track['name']), bold=True)
    bullet('养客：朋友圈/群科普建立"{}"认知，不硬推。'.format(track['awareness']))
    bullet('处关系：在群/朋友圈做捧场王，先建立信任。')
    bullet('挖需求：{}.'.format(track['dig']))
    bullet('邀客户：私聊 1v1，"{}"'.format(track['invite']))
    bullet('给方案：{}（不是单卖产品）。'.format(track['solution']))
    h2('第5步：表彰 / 毕业 —— 标杆上墙 + 目标分解')
    bullet('谁出单了/谁带出人了，群内"伙伴案例"上墙（模板D 周三）。', '【标杆】 ')
    bullet('目标分解 + 正向激励，让伙伴看见自己的成长曲线，而不是只看业绩数字。', '【激励】 ')

    # ---- 三、能力地图 ----
    h1('三、{}主理人能力地图（带教里程碑）'.format(track['name']))
    table(['阶段', '时间', '里程碑（做到这些就算过关）'],
          [['阶段1 上手', '第1-2周', '会用四件套工具，独立建一个{}打卡群并排出第一周内容'.format(track['name'])],
           ['阶段2 运营', '第3-4周', track['m2']],
           ['阶段3 出单', '第5-8周', track['m3']],
           ['阶段4 带人', '第9周+', '能带 1 个新人 = 出现潜在业务伙伴苗子']],
          widths=[1.0, 0.9, 4.3])
    para('带教节奏：每周看一个伙伴卡在哪个阶段，针对性给一个动作，不一次性塞一堆。', italic=True, color=AMBER)

    # ---- 四、周排期 ----
    h1('四、事业群一周排期（套模板D · {}主理人培养定制）'.format(track['name']))
    para('固定栏目固定时间（具体几点按伙伴作息定），写进群公告。每条带钩子+行动指令——内容离业务近。', bold=True)
    cols = ['周一','周二','周三','周四','周五','周六','周日']
    cols2 = ['事业科普','🔥带教时间','伙伴案例','品牌/产品知识','带教进阶','互动日','轻陪伴+预告']
    trust = ['二·认可你','二·认可你','三·信任你','二·认可你','二·认可你','一+三','一·认识你']
    rows = [[cols[i], cols2[i], track['week_ex'][i], trust[i], track['week_hook'][i]] for i in range(7)]
    table(['星期','固定栏目','内容示例（{}主理人培养）'.format(track['name']),'信任层','钩子+行动指令'], rows,
          widths=[0.5,0.95,2.5,0.7,1.55])

    # ---- 五、发现潜在业务伙伴 ----
    h1('五、发现潜在业务伙伴机制')
    para('事业群最核心的"产出"之一，是把优秀伙伴识别成未来的业务伙伴。', bold=True)
    bullet('KOC 培育：从{}打卡群筛选有意愿 + 有影响力的人，私聊问建议、给"内测资格"特殊感。'.format(track['name']), '① 筛选 ')
    bullet('让 TA 在群里带节奏（先晒打卡、先分享），你做捧场王+总结者，自然成为"自己人"。', '② 带节奏 ')
    bullet('一对一咨询钩子：识别出"想带人"信号的人，单独聊事业方向，不是群内喊。', '③ 深聊 ')
    bullet('信号清单：主动问"怎么带别人""这个能长期做吗""我想试试带群"= 高意向。', '④ 看信号 ')

    # ---- 六、销售&解决方案带教 ----
    h1('六、销售产品 & 解决方案带教（{}版）'.format(track['name']))
    para('教伙伴卖的是"方案"不是"货"——产品是方案的组成部分。', bold=True)
    table(['步骤','伙伴要会的话术/动作','带教要点'],
          [['养客', track['p0'], '不硬推，先给认知'],
           ['处关系', '在群/朋友圈做捧场王，接每一句', '信任先于成交'],
           ['挖需求', track['p2'], '用问卷/自测引出真实困扰'],
           ['邀客户', track['p3'], '从群转到私聊'],
           ['给方案', track['p4'], '方案>单品，讲清为什么']],
          widths=[0.8,3.0,2.4])

    # ---- 七、内容示例 ----
    h1('七、群内容示例（{}赛道 · 套四标签）'.format(track['track']))
    para('以下为事业群"带教时间"与"伙伴案例"可直接用的示例，已带钩子+行动指令。', bold=True)
    for ex in track['examples']:
        para(ex['title'], bold=True, color=AMBER)
        para(ex['body'])
        para('钩子：{} ｜ 行动指令：{}'.format(ex['hook'], ex['action']), italic=True, color=DARK)

    # ---- 八、合规底线 ----
    h1('八、合规底线（必守）')
    bullet('不夸大：不用"治愈/根除/最有效/第一"等绝对词。')
    bullet('不制造焦虑：鼓励接纳，不恐吓"再不调就晚了"。')
    bullet('建议咨询：涉及诊断/停药，引导找正规医生，不替代医疗建议。')

    # ---- 九、行动清单 ----
    h1('九、团队长本周就能做（行动清单）')
    bullet(track['action1'])
    bullet('把本周排期（第四节模板）写进群公告，定好固定栏目时间。')
    bullet('周二带教时间讲"{}破冰话术"，让伙伴现场练一句交作业。'.format(track['name']))
    bullet('建一个"伙伴成长"小表格，记录谁到哪个阶段、谁出单、谁有带人信号。')

    doc.add_paragraph()
    end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run('带出人，比卖出货更值钱。群会自己长。 —— 沁珮')
    er.bold = True; er.font.size = Pt(11); er.font.color.rgb = GREEN; set_cn(er)

    out = '{}/{}IP事业群带教方案.docx'.format(out_dir, track['name'])
    doc.save(out)
    print('SAVED:', out)

# ===================== 赛道参数（通用·不绑品牌） =====================
def get_tracks():
    tracks = []
    # 减脂
    tracks.append(dict(
        name='减脂', track='减脂塑形/体重管理',
        awareness='科学减脂认知',
        recruit='你这段时间打卡变化挺大，有没有想过把这套方法带给自己身边也卡在减重的人？我带你试一次。',
        dig='发身高体重自测，或体质小问卷，引出真实困扰',
        invite='"你这数据我帮你定制个方案？"1v1',
        solution='产品+饮食运动组合方案（如蛋白+膳食纤维+21天打卡）',
        m2='能独立做减脂科普+打卡运营，群活跃率≥20%',
        m3='能走通第一单：卖产品+给组合方案，敢做1v1咨询',
        week_ex=['为什么减脂赛道值得长期做（趋势+复购逻辑）','一个具体方法：减脂破冰话术怎么练','谁这周出单了/带出人了（上墙）','减脂产品专业度（蛋白/纤维/代谢）','话术复盘：伙伴卡点公开练','伙伴提问接龙/经验分享会','一周成长小结+下周预告'],
        week_hook=['"转给想了解的人"→引流意向','现场练一句→群内交作业','"你的故事上墙？"→征集标杆','懂产品才敢讲→小测巩固','抛出卡点→群内公开课','接龙提问→互助解答','一句鼓励→收尾不空'],
        p0='科普"科学减脂"认知建立', p2='发"身高体重"自测', p3='"你这数据我帮你定制个方案？"1v1', p4='产品+饮食运动组合',
        examples=[
            dict(title='【周二·带教时间】减脂破冰话术练习', body='很多伙伴跟客户开口第一句就卡住。试试这句："你坚持超3周还有效的减重方法是什么？"——先问再听，比直接推产品强10倍。现在把你想对老客户说的第一句打在群里，我帮你改。', hook='现场练一句', action='群内交作业'),
            dict(title='【周三·伙伴案例】主理人出单故事', body='上周@小雅 带出第一个减脂主理人：她陪伙伴用21天打卡+蛋白方案，伙伴自己瘦了4斤，转头带了2个朋友进群。这就是咱们要的——伙伴成长，群自己长。', hook='你的故事上墙？', action='征集本月标杆'),
            dict(title='【周四·产品知识】蛋白与减脂的关系', body='减脂不是不吃，是吃对。蛋白质够，才不掉肌肉、不反弹。很多伙伴瘦得快反弹也快，根因是蛋白没吃够。今天小测：一份早餐蛋白够不够？把你早餐发群里我帮你算。', hook='懂产品才敢讲', action='早餐小测巩固'),
        ],
        action1='从现有减脂打卡群/养客群里，挑 3 个"效果明显+爱分享"的人，私聊邀约进事业群。'))
    # 视康
    tracks.append(dict(
        name='视康', track='视康/青少年近视防控',
        awareness='近视不可逆但可防控',
        recruit='你这段时间护眼方法挺见效，有没有想过带给自己身边也愁孩子视力的朋友？我带你试一次。',
        dig='发"孩子读写姿势/用眼时长/户外时长"自测，引出真实困扰',
        invite='"你这数据我帮你判断风险，私聊给你一份家庭护眼建议"',
        solution='护眼营养（叶黄素类）+ 用眼习惯方案 + 家庭护眼打卡',
        m2='能独立做护眼科普+打卡运营，群活跃率≥20%', m3='能走通第一单：卖护眼组合方案，敢做1v1咨询',
        week_ex=['为什么视康赛道值得长期做（防控市场+复购逻辑）','一个具体方法：护眼破冰话术怎么练','谁这周出单了/带出人了（上墙）','护眼产品专业度（叶黄素/蓝光/眼肌训练）','话术复盘：伙伴卡点公开练','伙伴提问接龙/经验分享会','一周成长小结+下周预告'],
        week_hook=['"转给想了解的人"→引流意向','现场练一句→群内交作业','"你的故事上墙？"→征集标杆','懂产品才敢讲→小测巩固','抛出卡点→群内公开课','接龙提问→互助解答','一句鼓励→收尾不空'],
        p0='科普"近视可防控"认知建立', p2='发"读写姿势/户外时长"自测', p3='"你这数据我帮你判断风险？"1v1', p4='护眼营养+用眼习惯+家庭打卡组合',
        examples=[
            dict(title='【周二·带教时间】护眼破冰话术练习', body='很多伙伴跟家长开口第一句就卡住。试试这句："你家孩子每天户外够1小时吗？"——先问再听，比直接推产品强10倍。现在把你想对老客户说的第一句打在群里，我帮你改。', hook='现场练一句', action='群内交作业'),
            dict(title='【周三·伙伴案例】主理人出单故事', body='上周@小雅 带出第一个护眼主理人：她陪伙伴用家庭护眼打卡+叶黄素方案，伙伴自己更注意用眼了，转头带了2个朋友进群。这就是咱们要的——伙伴成长，群自己长。', hook='你的故事上墙？', action='征集本月标杆'),
            dict(title='【周四·产品知识】叶黄素与眼健康', body='护眼不是少看手机就行，是营养+习惯。叶黄素是眼底黄斑的"防晒伞"，很多孩子挑食、户外少，储备不够。今天小测：孩子今天户外多久？把你家数据发群里我帮你算。', hook='懂产品才敢讲', action='户外小测巩固'),
        ],
        action1='从现有护眼打卡群/养客群里，挑 3 个"效果明显+爱分享"的家长，私聊邀约进事业群。'))
    # 艾灸
    tracks.append(dict(
        name='艾灸', track='艾灸/中医理疗',
        awareness='节气艾灸/寒湿调理',
        recruit='你这段时间灸完舒服多了，有没有想过带给自己身边也怕冷的姐妹？我带你试一次。',
        dig='发"症状自测（怕冷/乏力/睡不好/痛经）"，引出真实体质困扰',
        invite='"发症状，抽3位1对1辨证，私聊给你调理建议"',
        solution='艾灸产品+一茶一灸+体质调理方案',
        m2='能独立做艾灸科普+打卡运营，群活跃率≥20%', m3='能走通第一单：卖艾灸+调理方案，敢做1v1辨证',
        week_ex=['为什么艾灸/中医外治赛道值得长期做（刚需+复购）','一个具体方法：艾灸破冰话术怎么练','谁这周出单了/带出人了（上墙）','艾灸产品专业度（灸具/艾条/配穴/体质）','话术复盘：伙伴卡点公开练','伙伴提问接龙/经验分享会','一周成长小结+下周预告'],
        week_hook=['"转给想了解的人"→引流意向','现场练一句→群内交作业','"你的故事上墙？"→征集标杆','懂产品才敢讲→小测巩固','抛出卡点→群内公开课','接龙提问→互助解答','一句鼓励→收尾不空'],
        p0='科普"节气艾灸/寒湿调理"认知', p2='发"症状自测（怕冷/乏力/睡不好）"', p3='"发症状，抽3位1对1辨证"1v1', p4='艾灸+一茶一灸+体质调理组合',
        examples=[
            dict(title='【周二·带教时间】艾灸破冰话术练习', body='很多伙伴跟姐妹开口第一句就卡住。试试这句："你平时哪里最怕冷？"——先问再听，比直接推灸具强10倍。现在把你想对老客户说的第一句打在群里，我帮你改。', hook='现场练一句', action='群内交作业'),
            dict(title='【周三·伙伴案例】主理人出单故事', body='上周@王姐 带出第一个艾灸主理人：她陪伙伴用节气艾灸+姜枣茶方案，伙伴自己不怕冷了，转头带了2个闺蜜进群。这就是咱们要的——伙伴成长，群自己长。', hook='你的故事上墙？', action='征集本月标杆'),
            dict(title='【周四·产品知识】艾与寒湿体质', body='艾灸不是哪疼灸哪，是辨证。寒湿体质的人灸完排寒反应大，反而怕。今天小测：你是什么体质？把怕冷/乏力情况发群里我帮你辨。', hook='懂产品才敢讲', action='体质小测巩固'),
        ],
        action1='从现有艾灸打卡群/养客群里，挑 3 个"改善明显+爱分享"的姐妹，私聊邀约进事业群。'))
    # 儿童喂养
    tracks.append(dict(
        name='儿童喂养', track='儿童健康喂养/脾胃调理',
        awareness='科学喂养/脾胃调理',
        recruit='你孩子这段时间吃饭香了，有没有想过带给自己身边也愁喂养的朋友？我带你试一次。',
        dig='发"孩子身高体重/挑食情况/作息"自测，引出真实喂养困扰',
        invite='"发数据，私聊给你一份专属喂养方案"',
        solution='儿童营养品+喂养作息+推拿/食疗方案',
        m2='能独立做喂养科普+打卡运营，群活跃率≥20%', m3='能走通第一单：卖儿童营养+喂养方案，敢做1v1咨询',
        week_ex=['为什么儿童喂养赛道值得长期做（焦虑刚需+转介绍）','一个具体方法：喂养破冰话术怎么练','谁这周出单了/带出人了（上墙）','儿童营养专业度（脾胃/吸收/关键营养素）','话术复盘：伙伴卡点公开练','伙伴提问接龙/经验分享会','一周成长小结+下周预告'],
        week_hook=['"转给想了解的人"→引流意向','现场练一句→群内交作业','"你的故事上墙？"→征集标杆','懂产品才敢讲→小测巩固','抛出卡点→群内公开课','接龙提问→互助解答','一句鼓励→收尾不空'],
        p0='科普"科学喂养/脾胃调理"认知', p2='发"身高体重/挑食情况"自测', p3='"发数据，私你喂养方案"1v1', p4='儿童营养+喂养作息+推拿食疗组合',
        examples=[
            dict(title='【周二·带教时间】喂养破冰话术练习', body='很多伙伴跟妈妈开口第一句就卡住。试试这句："孩子挑食你硬塞还是换花样？"——先问再听，比直接推营养品强10倍。现在把你想对老客户说的第一句打在群里，我帮你改。', hook='现场练一句', action='群内交作业'),
            dict(title='【周三·伙伴案例】主理人出单故事', body='上周@乐乐妈 带出第一个喂养主理人：她陪伙伴用喂养作息+儿童营养方案，孩子吃饭香了，转头带了2个同学妈妈进群。这就是咱们要的——伙伴成长，群自己长。', hook='你的故事上墙？', action='征集本月标杆'),
            dict(title='【周四·产品知识】脾胃与营养吸收', body='孩子不是吃得少，是吸收差。脾胃弱，补再多也白搭。今天小测：孩子早餐吃对没？把早餐发群里我帮你算吸收。', hook='懂产品才敢讲', action='早餐小测巩固'),
        ],
        action1='从现有喂养打卡群/养客群里，挑 3 个"孩子变化明显+爱分享"的家长，私聊邀约进事业群。'))
    # 40+体质调理
    tracks.append(dict(
        name='40+体质调理', track='40+体质调理/中医食疗',
        awareness='九种体质/节气食疗',
        recruit='你这段时间气色好多了，有没有想过带给自己身边也累的朋友？我带你试一次。',
        dig='发"体质自测（气血/湿寒/虚）"，引出真实亚健康困扰',
        invite='"做测试领食谱，私聊给你体质调理建议"',
        solution='食疗+外养（灸/饮）+体质调理方案',
        m2='能独立做体质科普+打卡运营，群活跃率≥20%', m3='能走通第一单：卖体质调理组合方案，敢做1v1咨询',
        week_ex=['为什么40+体质调理赛道值得长期做（刚需+高客单）','一个具体方法：体质破冰话术怎么练','谁这周出单了/带出人了（上墙）','体质调理专业度（九种体质/节气食疗/外养）','话术复盘：伙伴卡点公开练','伙伴提问接龙/经验分享会','一周成长小结+下周预告'],
        week_hook=['"转给想了解的人"→引流意向','现场练一句→群内交作业','"你的故事上墙？"→征集标杆','懂产品才敢讲→小测巩固','抛出卡点→群内公开课','接龙提问→互助解答','一句鼓励→收尾不空'],
        p0='科普"九种体质/节气食疗"认知', p2='发"体质自测（气血/湿寒/虚）"', p3='"做测试领食谱"1v1', p4='食疗+外养+体质调理组合',
        examples=[
            dict(title='【周二·带教时间】体质破冰话术练习', body='很多伙伴跟姐妹开口第一句就卡住。试试这句："总犯困气血虚？你平时怎么调？"——先问再听，比直接推产品强10倍。现在把你想对老客户说的第一句打在群里，我帮你改。', hook='现场练一句', action='群内交作业'),
            dict(title='【周三·伙伴案例】主理人出单故事', body='上周@丽姐 带出第一个体质调理主理人：她陪伙伴用节气食疗+温养打卡方案，自己气色好了，转头带了2个闺蜜进群。这就是咱们要的——伙伴成长，群自己长。', hook='你的故事上墙？', action='征集本月标杆'),
            dict(title='【周四·产品知识】体质与食疗', body='40+不是补越多越好，是辨体质。湿寒质猛补上火，气虚质虚不受补。今天小测：你是哪种体质？把症状发群里我帮你辨。', hook='懂产品才敢讲', action='体质小测巩固'),
        ],
        action1='从现有体质打卡群/养客群里，挑 3 个"气色变好+爱分享"的姐妹，私聊邀约进事业群。'))
    return tracks

if __name__ == '__main__':
    import os
    out_dir = sys.argv[1] if len(sys.argv) > 1 else '/Users/queen1015/Desktop'
    # 支持只生成指定赛道：python generate_plan.py <dir> 减脂 视康
    want = sys.argv[2:] if len(sys.argv) > 2 else None
    for t in get_tracks():
        if want and t['name'] not in want:
            continue
        build(t, out_dir)
    print('ALL DONE')
